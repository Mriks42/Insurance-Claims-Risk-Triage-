"""Generate the Status Update 2 Quad Chart as a Word document."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

# ── Styles helper ──
def add_title(text):
    p = doc.add_heading(text, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0, 51, 102)

def add_subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(100, 100, 100)

def add_quadrant_heading(text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 82, 136)

def add_quadrant_subtitle(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(60, 60, 60)
    run.font.italic = True

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run_bold = p.add_run(bold_prefix)
        run_bold.bold = True
        run_bold.font.size = Pt(10)
        run_rest = p.add_run(text)
        run_rest.font.size = Pt(10)
    else:
        p.runs[0].font.size = Pt(10) if p.runs else None
        p.clear()
        run = p.add_run(text)
        run.font.size = Pt(10)

def add_sub_bullet(text):
    p = doc.add_paragraph(style="List Bullet 2")
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(10)

def add_separator():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("─" * 80)
    run.font.size = Pt(6)
    run.font.color.rgb = RGBColor(180, 180, 180)

# ══════════════════════════════════════════════
# DOCUMENT CONTENT
# ══════════════════════════════════════════════

add_title("Status Update 2 — Quad Chart")
add_subtitle("FSE 570 Data Science Capstone  |  Team Connecticut")
add_subtitle("Mriganko Chowdhury  •  Aryan Gonsalves  •  Ashish Raj Singh  •  Deborah Sheryl Veluvalli  •  Kshama Girish")

add_separator()

# ── Q1 ──
add_quadrant_heading("Q1: Project Scope & Status")
add_quadrant_subtitle("Phases 3–5 Complete | On Track for Final Demo")

add_bullet("Transitioned from Phase 3 (Advanced Modeling) through Phase 5 (RAG Pipeline) in this reporting period")
add_bullet("Modular Refactoring — ", "Modular Refactoring — ")
doc.paragraphs[-1].clear()
p = doc.paragraphs[-1]
run = p.add_run("Modular Refactoring — ")
run.bold = True
run.font.size = Pt(10)
run = p.add_run("Migrated 60-cell monolithic notebook into 8 importable Python modules with a unified outputs/ directory (100% Complete)")
run.font.size = Pt(10)

add_bullet("", "Data Leakage Audit — ")
p = doc.paragraphs[-1]
p.clear()
run = p.add_run("Data Leakage Audit — ")
run.bold = True
run.font.size = Pt(10)
run = p.add_run("Identified and removed PolicyNumber identifier from the feature set; documented impact on all downstream metrics (100% Complete)")
run.font.size = Pt(10)
p.style = doc.styles["List Bullet"]

add_bullet("", "SHAP Explainability — ")
p = doc.paragraphs[-1]
p.clear()
run = p.add_run("SHAP Explainability — ")
run.bold = True
run.font.size = Pt(10)
run = p.add_run("Integrated TreeExplainer for global feature importance and per-claim human-readable reason codes for all SIU-tier claims (100% Complete)")
run.font.size = Pt(10)
p.style = doc.styles["List Bullet"]

add_bullet("", "Feature Engineering & Model Improvement — ")
p = doc.paragraphs[-1]
p.clear()
run = p.add_run("Feature Engineering & Model Improvement — ")
run.bold = True
run.font.size = Pt(10)
run = p.add_run("Built feature_engineering.py (34 new features) and model_improvement.py (hyperparameter tuning via RandomizedSearchCV + CatBoost comparison), yielding a +25% PR-AUC improvement (100% Complete)")
run.font.size = Pt(10)
p.style = doc.styles["List Bullet"]

add_bullet("", "RAG Pipeline with Fallback — ")
p = doc.paragraphs[-1]
p.clear()
run = p.add_run("RAG Pipeline with Fallback — ")
run.bold = True
run.font.size = Pt(10)
run = p.add_run("Built end-to-end retrieval-augmented generation system with ChromaDB vector index, 3 curated guideline documents, and a template-based fallback that operates without an API key (100% Complete)")
run.font.size = Pt(10)
p.style = doc.styles["List Bullet"]

add_bullet("Overall progress: 5 of 6 phases complete; Phase 6 (Streamlit Dashboard) is the remaining deliverable")

add_separator()

# ── Q2 ──
add_quadrant_heading("Q2: Systems Architecture & Engineering")
add_quadrant_subtitle("Modular Pipeline + Local-First RAG Engine")

add_bullet("", "Modular Architecture (8 modules): ")
p = doc.paragraphs[-1]
p.clear()
run = p.add_run("Modular Architecture (8 modules): ")
run.bold = True
run.font.size = Pt(10)
run = p.add_run("config.py → data_preprocessing.py → modeling.py → feature_engineering.py → model_improvement.py → shap_explainability.py → rag_pipeline.py — each is independently runnable and importable")
run.font.size = Pt(10)
p.style = doc.styles["List Bullet"]

add_bullet("", "Preprocessing: ")
p = doc.paragraphs[-1]
p.clear()
run = p.add_run("Preprocessing: ")
run.bold = True
run.font.size = Pt(10)
run = p.add_run("Stratified 80/10/10 split, ColumnTransformer with median imputation + StandardScaler (numeric) and OneHotEncoder (categorical); 147 transformed features")
run.font.size = Pt(10)
p.style = doc.styles["List Bullet"]

add_bullet("", "Model Suite (modeling.py): ")
p = doc.paragraphs[-1]
p.clear()
run = p.add_run("Model Suite (modeling.py): ")
run.bold = True
run.font.size = Pt(10)
run = p.add_run("Logistic Regression baseline, XGBoost (best), LightGBM — all evaluated with capacity-aware metrics (PR-AUC, Precision@K, Recall@K)")
run.font.size = Pt(10)
p.style = doc.styles["List Bullet"]

add_bullet("", "Feature Engineering (feature_engineering.py): ")
p = doc.paragraphs[-1]
p.clear()
run = p.add_run("Feature Engineering (feature_engineering.py): ")
run.bold = True
run.font.size = Pt(10)
run = p.add_run("34 engineered features — ordinal conversions (9 columns), time gaps (3), binary risk flags (9), interaction terms (6), and a composite RiskFlagCount")
run.font.size = Pt(10)
p.style = doc.styles["List Bullet"]

add_bullet("", "Model Improvement (model_improvement.py): ")
p = doc.paragraphs[-1]
p.clear()
run = p.add_run("Model Improvement (model_improvement.py): ")
run.bold = True
run.font.size = Pt(10)
run = p.add_run("XGBoost with engineered features, RandomizedSearchCV hyperparameter tuning (40 iterations, 3-fold CV), and CatBoost (native categorical handling) — all three compared side-by-side")
run.font.size = Pt(10)
p.style = doc.styles["List Bullet"]

add_bullet("", "RAG Engine (rag_pipeline.py):")
p = doc.paragraphs[-1]
p.clear()
run = p.add_run("RAG Engine (rag_pipeline.py):")
run.bold = True
run.font.size = Pt(10)
p.style = doc.styles["List Bullet"]

add_sub_bullet("Embedding: all-MiniLM-L6-v2 sentence-transformer (runs locally, no API required)")
add_sub_bullet("Vector Store: ChromaDB in-memory index with 36 semantic chunks from 3 guideline documents (~11,600 chars)")
add_sub_bullet("Retrieval: Claim data + SHAP reason codes → natural language query → top-5 passage retrieval")
add_sub_bullet("Generation: Dual-mode — OpenAI GPT-4o-mini (if API key available) OR structured template fallback with cited policy passages")

add_bullet("", "Reproducibility: ")
p = doc.paragraphs[-1]
p.clear()
run = p.add_run("Reproducibility: ")
run.bold = True
run.font.size = Pt(10)
run = p.add_run("Single command (python shap_explainability.py) runs the full pipeline end-to-end; all artifacts saved to outputs/")
run.font.size = Pt(10)
p.style = doc.styles["List Bullet"]

add_separator()

# ── Q3 ──
add_quadrant_heading("Q3: Critical Results & Integrity Victory")
add_quadrant_subtitle("Honest Metrics + 25% Recovery + 4.5x Triage Power")

add_bullet("", "Data Integrity Victory: ")
p = doc.paragraphs[-1]
p.clear()
run = p.add_run("Data Integrity Victory: ")
run.bold = True
run.font.size = Pt(10)
run = p.add_run("Removing the PolicyNumber identifier (a row-level unique key leaked into features) dropped Test PR-AUC from 0.79 → 0.19 — confirming the original model was memorizing IDs, not learning fraud patterns. The corrected model produces honest, production-representative metrics")
run.font.size = Pt(10)
p.style = doc.styles["List Bullet"]

add_bullet("", "Feature Engineering Recovery (feature_engineering.py): ")
p = doc.paragraphs[-1]
p.clear()
run = p.add_run("Feature Engineering Recovery (feature_engineering.py): ")
run.bold = True
run.font.size = Pt(10)
run = p.add_run("34 engineered features boosted Val PR-AUC from 0.2518 → 0.3149 (+25.1%) and Test PR-AUC from 0.1951 → 0.2586 (+32.5%) — the single most impactful improvement")
run.font.size = Pt(10)
p.style = doc.styles["List Bullet"]

add_bullet("", "Model Improvement Comparison (model_improvement.py): ")
p = doc.paragraphs[-1]
p.clear()
run = p.add_run("Model Improvement Comparison (model_improvement.py): ")
run.bold = True
run.font.size = Pt(10)
run = p.add_run("XGBoost+FeatEng (PR-AUC 0.3149) > CatBoost (0.3048) > XGBoost Tuned (0.2997) — feature engineering outperformed both hyperparameter tuning and an entirely different algorithm")
run.font.size = Pt(10)
p.style = doc.styles["List Bullet"]

add_bullet("", "Triage Enrichment:")
p = doc.paragraphs[-1]
p.clear()
run = p.add_run("Triage Enrichment:")
run.bold = True
run.font.size = Pt(10)
p.style = doc.styles["List Bullet"]

add_sub_bullet("SIU bucket (top 5%): 27.3% fraud rate → 4.5x enrichment vs. base rate (6.0%)")
add_sub_bullet("Manual Review (next 15%): 18.6% fraud rate → 3.1x enrichment")
add_sub_bullet("Approve (remaining 80%): 2.4% fraud rate → fraud leakage minimized")

add_bullet("", "Top Fraud Drivers (SHAP): ")
p = doc.paragraphs[-1]
p.clear()
run = p.add_run("Top Fraud Drivers (SHAP): ")
run.bold = True
run.font.size = Pt(10)
run = p.add_run("BasePolicy=Liability (0.108), Fault=PolicyHolder (0.096), PolicyType=Sedan-Collision (0.014), Age (0.011), AddressChange=2-3yrs (0.008)")
run.font.size = Pt(10)
p.style = doc.styles["List Bullet"]

add_bullet("", "RAG Output: ")
p = doc.paragraphs[-1]
p.clear()
run = p.add_run("RAG Output: ")
run.bold = True
run.font.size = Pt(10)
run = p.add_run("All 3 triage tiers generate cited briefs with guideline references; SIU briefs include 5-step investigation protocols sourced from triage_procedures.md")
run.font.size = Pt(10)
p.style = doc.styles["List Bullet"]

add_separator()

# ── Q4 ──
add_quadrant_heading("Q4: Road Ahead & Next Steps")
add_quadrant_subtitle("Phase 6 — Streamlit Dashboard (Weeks 11–12)")

add_bullet("", "Goal: ")
p = doc.paragraphs[-1]
p.clear()
run = p.add_run("Goal: ")
run.bold = True
run.font.size = Pt(10)
run = p.add_run("Interactive web application for claims adjusters and SIU analysts to review flagged claims with full model transparency")
run.font.size = Pt(10)
p.style = doc.styles["List Bullet"]

add_bullet("", "Risk-Ranked Review Queue: ")
p = doc.paragraphs[-1]
p.clear()
run = p.add_run("Risk-Ranked Review Queue: ")
run.bold = True
run.font.size = Pt(10)
run = p.add_run("Sortable, filterable table of all claims with risk scores, triage bucket assignments, and fraud indicators — adjusters see highest-risk claims first")
run.font.size = Pt(10)
p.style = doc.styles["List Bullet"]

add_bullet("", "Per-Claim Detail View:")
p = doc.paragraphs[-1]
p.clear()
run = p.add_run("Per-Claim Detail View:")
run.bold = True
run.font.size = Pt(10)
p.style = doc.styles["List Bullet"]

add_sub_bullet("Fraud risk score with severity badge (HIGH / ELEVATED / MODERATE / LOW)")
add_sub_bullet("SHAP reason codes (top 5 risk drivers in plain language)")
add_sub_bullet("RAG-generated triage brief with cited guideline passages")
add_sub_bullet("Recommended investigation steps (SIU protocol or review checklist)")

add_bullet("", "Summary Dashboard: ")
p = doc.paragraphs[-1]
p.clear()
run = p.add_run("Summary Dashboard: ")
run.bold = True
run.font.size = Pt(10)
run = p.add_run("Aggregate visualizations including fraud rate by triage bucket, global SHAP feature importance chart, Precision@K curves, and model comparison metrics")
run.font.size = Pt(10)
p.style = doc.styles["List Bullet"]

add_bullet("", "Technical Integration: ")
p = doc.paragraphs[-1]
p.clear()
run = p.add_run("Technical Integration: ")
run.bold = True
run.font.size = Pt(10)
run = p.add_run("All modules already expose importable functions; the Streamlit app will import from data_preprocessing, modeling, shap_explainability, and rag_pipeline directly")
run.font.size = Pt(10)
p.style = doc.styles["List Bullet"]

add_bullet("", "Deliverable Timeline: ")
p = doc.paragraphs[-1]
p.clear()
run = p.add_run("Deliverable Timeline: ")
run.bold = True
run.font.size = Pt(10)
run = p.add_run("Functional prototype by Week 11; polish, testing, and final demo preparation in Week 12")
run.font.size = Pt(10)
p.style = doc.styles["List Bullet"]

# ── Save ──
out_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Status_Update_2_Quad_Chart.docx")
doc.save(out_path)
print(f"Saved: {out_path}")
